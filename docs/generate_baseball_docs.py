"""Generate Word documents for baseball model summaries."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DIR = os.path.dirname(os.path.abspath(__file__))

for md_file, title, subtitle in [
    ("BBMI_Baseball_Model_Internal_Summary.md", "BBMI Baseball Model (V2)", "Internal Technical Summary"),
    ("BBMI_Baseball_Model_Review_Brief.md", "BBMI Sports - NCAA Baseball Predictive Model", "Methodology Brief for External Review"),
]:
    path = os.path.join(DIR, md_file)
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("March 2026 | CONFIDENTIAL")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)

    doc.add_page_break()

    for line in content.split("\n"):
        line = line.rstrip()
        if line.startswith("# ") and not line.startswith("##"):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("| ") and "---" not in line:
            p = doc.add_paragraph(line)
            p.style = "No Spacing"
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Consolas"
        elif line.startswith("```"):
            continue
        elif line.startswith("- **") or line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("---"):
            continue
        elif line.strip():
            doc.add_paragraph(line)

    out_path = os.path.join(DIR, md_file.replace(".md", ".docx"))
    doc.save(out_path)
    print(f"Created: {out_path}")
