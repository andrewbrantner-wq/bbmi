"""Generate Word documents for BBMI Football Model summaries."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

NAVY = RGBColor(0x0A, 0x16, 0x28)
DARK_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GRAY = RGBColor(0x78, 0x71, 0x6C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, header_bg="0A1628", col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = WHITE
        set_cell_shading(cell, header_bg)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
            if r % 2 == 0:
                set_cell_shading(cell, "F5F5F4")

    return table


def build_internal_doc():
    """Build the internal technical summary."""
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BBMI Football Model")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Internal Technical Summary")
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("March 2026 | CONFIDENTIAL")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True

    doc.add_paragraph()

    # Overview
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "The BBMI Football model (BBMIF) is a NCAA Division I football predictive system that generates "
        "spread predictions, over/under totals, and win probabilities for every FBS game. It powers the "
        "bbmisports.com picks, rankings, and accuracy pages."
    )

    # Data Sources
    doc.add_heading("Data Sources", level=1)
    add_styled_table(doc,
        ["Source", "Data", "Auth", "Refresh"],
        [
            ["ESPN API", "Schedules, scores, box scores (yards, turnovers, 3rd down, sacks)", "Free", "Daily, 3-day lookback"],
            ["CFBD API", "SP+ ratings (offense, defense, special teams), season stats", "Free (API key)", "Every 6 days"],
            ["The Odds API", "Vegas spreads, totals, moneylines (DK/FD/BetMGM priority)", "Paid", "Each pipeline run"],
            ["Open-Meteo", "Historical weather (temp, wind, precipitation)", "Free", "Cached per venue/date"],
        ]
    )

    # Stat Weights
    doc.add_heading("Rating System: Stat Weights", level=1)
    doc.add_paragraph(
        "Each FBS team is percentile-ranked 1-50 on each stat, then combined via weighted sum. "
        "SP+ carries 76% of model weight because it is opponent-adjusted and the most predictive "
        "single input available for college football. Box score stats serve as a recency signal."
    )
    add_styled_table(doc,
        ["Stat", "Weight", "Source", "Notes"],
        [
            ["SP+ Offense", "38.0%", "CFBD", "Opponent-adjusted offensive efficiency"],
            ["SP+ Defense", "38.0%", "CFBD", "Opponent-adjusted defensive efficiency"],
            ["Turnover Margin", "4.0%", "ESPN box scores", "Turnovers forced - committed per game"],
            ["3rd Down Diff", "4.0%", "ESPN box scores", "Conversion rate differential"],
            ["Yards Per Play Diff", "3.0%", "ESPN box scores", "Net yards per play advantage"],
            ["SP+ Special Teams", "3.0%", "CFBD", "Kick/punt return efficiency"],
            ["Quality Wins", "~3.0%", "Calculated", "Tier-weighted wins vs ranked opponents"],
            ["Last 5 Record", "2.0%", "Calculated", "Wins in last 5 games (0-5)"],
        ]
    )

    # BBMIF Score
    doc.add_heading("BBMIF Score Calculation", level=2)
    doc.add_paragraph("raw_bbmif = sum(percentile_rank[stat] * weight[stat]) * 0.8", style='No Spacing')
    doc.add_paragraph()
    doc.add_paragraph("Rank caps are applied iteratively:")
    doc.add_paragraph("Best-win cap: Cannot rank more than 15 positions above best-beaten opponent's rank", style='List Bullet')
    doc.add_paragraph("Day-over-day cap: Cannot improve more than 12 positions from yesterday", style='List Bullet')

    # Quality Wins
    doc.add_heading("Quality Wins Scoring", level=2)
    add_styled_table(doc,
        ["Opponent Tier", "Win", "Close Loss (<=3 pts)", "Loss"],
        [
            ["Top 10", "+6.0", "+1.5", "-1.5"],
            ["11-20", "+3.0", "+0.5", "-3.0"],
            ["21-25", "+3.0", "--", "-4.5"],
            ["26-40", "+1.0", "--", "--"],
            ["41-75", "+0.4", "--", "--"],
        ]
    )

    # Spread Prediction
    doc.add_heading("Spread Prediction", level=1)
    doc.add_heading("Formula", level=2)
    p = doc.add_paragraph()
    run = p.add_run("raw_line = (away_bbmif - home_bbmif + HFA) * 1.68\nfinal_line = snap_to_half_point(raw_line + 1.0)")
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading("Key Constants", level=2)
    add_styled_table(doc,
        ["Parameter", "Value", "Calibration Method"],
        [
            ["BBMIF_LINE_MULTIPLIER", "1.68", "OLS regression on 796 games"],
            ["STD_DEV", "14.0", "Brier score minimization"],
            ["HOME_BIAS_CORRECTION", "+1.0 pts", "Corrects home over-selection"],
            ["BYE_WEEK_BONUS", "2.5 pts", "Applied to rested team"],
            ["HCA_SCALE", "0.5", "Scales raw residual-based HFA"],
            ["HCA_CAP", "+/-8.0 pts", "Prevents extreme HFA values"],
        ]
    )

    doc.add_heading("Home Field Advantage (Residual Method)", level=2)
    doc.add_paragraph(
        "HFA is calculated per team from actual game margins vs projected margins: "
        "true_HFA = (avg_home_residual - avg_away_residual) / 2, then scaled by 0.5 and capped at +/-8.0 points."
    )

    doc.add_heading("Altitude Adjustments", level=2)
    add_styled_table(doc,
        ["Team", "Bonus (pts)"],
        [
            ["Air Force", "4.0"], ["Wyoming", "3.0"], ["Colorado / Colorado State", "2.5"],
            ["BYU / Utah", "2.0"], ["New Mexico", "1.5"], ["UNLV", "2.0"],
        ]
    )

    # Over/Under
    doc.add_heading("Over/Under Calculation", level=1)
    doc.add_heading("Base Point Projections (SP+ Method)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "team_pts = regressed_offense * (29.0 / opponent_regressed_defense)\n"
        "where regressed = raw * 0.80 + 29.0 * 0.20  (20% regression to mean)"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading("Conference Tier Total Offsets", level=2)
    add_styled_table(doc,
        ["Matchup", "Offset"],
        [
            ["Power 4 vs Power 4", "-23.6 pts"],
            ["Power 4 vs Group of 5", "-12.6 pts"],
            ["Group of 5 vs Group of 5", "+4.7 pts"],
        ]
    )

    doc.add_heading("Weather Adjustments (outdoor games)", level=2)
    add_styled_table(doc,
        ["Condition", "Impact"],
        [
            ["Wind > 20 mph", "-4.0 pts"],
            ["Wind 15-20 mph", "-2.0 pts"],
            ["Temperature < 20F", "-3.0 pts"],
            ["Temperature 20-32F", "-1.5 pts"],
            ["Precipitation > 2mm", "-2.0 pts"],
        ]
    )

    doc.add_heading("Pace Adjustment", level=2)
    doc.add_paragraph(
        "Expected plays per game (blending team offensive pace with opponent defensive pace) "
        "adjusts total by +/-20% relative to the 70-play FBS average."
    )

    # Win Probability
    doc.add_heading("Win Probability", level=1)
    doc.add_paragraph(
        "Uses Gaussian CDF (complementary error function) with the predicted spread as the mean "
        "and STD_DEV = 14.0. The actual observed standard deviation is 19.8, but 14.0 produces "
        "better-calibrated probabilities via Brier score minimization."
    )
    p = doc.add_paragraph()
    run = p.add_run("home_win_prob = normal_cdf(spread, mean=0, std_dev=14.0)")
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Recency
    doc.add_heading("Recency & Seasonal Adjustments", level=1)
    add_styled_table(doc,
        ["Parameter", "Value", "Purpose"],
        [
            ["RECENCY_DECAY", "0.85/week", "Exponential decay weighting recent games higher"],
            ["EARLY_SEASON_PENALTY", "0.75", "Reduce box score influence weeks 1-3"],
            ["MIN_GAMES_FOR_BOXSCORE", "7", "Fall back to season stats if fewer games available"],
        ]
    )

    # Known Issues
    doc.add_heading("Known Issues", level=1)
    issues = [
        ("SP+ Dependency (76%)", "If SP+ is wrong about a team, BBMIF will be wrong."),
        ("Home Team Over-Selection", "Model picks home teams 70.5% vs 59.6% actual home win rate."),
        ("Mid-Season Accuracy Drop", "Weeks 1-3: 83% ATS. Weeks 4-9: 52% ATS. Box score noise overwhelms SP+ signal."),
        ("No Injury Modeling", "Key QB absences can swing 7-14 points. Not accounted for."),
        ("No Coaching Adjustments", "First-year coaches, coordinator changes, scheme shifts unmodeled."),
        ("Static Conference Offsets", "Total offsets don't adjust for late-season form changes."),
    ]
    for title, desc in issues:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # Pipeline
    doc.add_heading("Pipeline Architecture", level=1)
    doc.add_paragraph(
        "Pipeline: bbmi_football_pipeline.py (~5,500 lines)\n"
        "Inputs: ESPN API, CFBD API, Odds API, Open-Meteo\n"
        "Outputs: football-games.json, football-rankings.json\n"
        "Calibration: weight grid search, multiplier OLS, STD_DEV Brier, total offset bias, rolling backtest"
    )

    return doc


def build_review_doc():
    """Build the external review brief for the Vegas director."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Title page
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ANALYTICS CONSULTING REVIEW")
    run.font.size = Pt(10)
    run.font.color.rgb = GOLD
    run.bold = True
    run.font.letter_spacing = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BBMI Sports - NCAA Football Predictive Model")
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Methodology Brief for External Review")
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    run.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for: Director of Risk, Vegas Analytics\nMarch 2026 | CONFIDENTIAL")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # Purpose
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "BBMI Sports operates a public-facing NCAA football predictive model at bbmisports.com that generates "
        "game-by-game spread predictions, over/under totals, and win probabilities for all FBS matchups. "
        "We are seeking a professional critique of our methodology from a Vegas pricing perspective to identify "
        "modeling errors, missing factors, and opportunities to improve line accuracy."
    )
    doc.add_paragraph(
        "We recently completed this process for our NCAA baseball model and implemented 17 recommended improvements. "
        "We are now turning the same lens on our football model ahead of the 2026 season."
    )

    # Model Architecture
    doc.add_heading("Model Architecture", level=1)
    doc.add_heading("Rating System", level=2)
    doc.add_paragraph(
        "Each FBS team receives a composite BBMIF score derived from a weighted blend of opponent-adjusted "
        "power ratings and box score performance metrics. Teams are percentile-ranked 1-50 on each input, "
        "then combined via weighted sum."
    )

    doc.add_heading("Input Weights", level=3)
    add_styled_table(doc,
        ["Input", "Weight", "Source", "Description"],
        [
            ["SP+ Offense", "38%", "CFBD (Bill Connelly)", "Opponent-adjusted per-play offensive efficiency"],
            ["SP+ Defense", "38%", "CFBD (Bill Connelly)", "Opponent-adjusted per-play defensive efficiency"],
            ["Turnover Margin", "4%", "ESPN box scores", "Turnovers forced minus committed per game"],
            ["3rd Down Differential", "4%", "ESPN box scores", "Conversion rate differential"],
            ["Yards Per Play Diff", "3%", "ESPN box scores", "Net yards per play advantage"],
            ["SP+ Special Teams", "3%", "CFBD", "Kick/punt return efficiency rating"],
            ["Quality Wins", "~3%", "Calculated", "Tier-weighted wins vs ranked opponents"],
            ["Last 5 Record", "2%", "Calculated", "Simple momentum signal"],
        ]
    )

    doc.add_paragraph(
        "Design rationale: SP+ carries 76% of model weight because it is opponent-adjusted and the most predictive "
        "single input available for college football. Box score stats are opponent-unadjusted and primarily serve as "
        "a recency signal once 7+ games are available. Before week 4, we apply an early-season penalty (0.75x) to "
        "box score weights and rely more heavily on SP+ preseason projections."
    )

    # Quality Wins
    doc.add_heading("Quality Wins Scoring", level=2)
    doc.add_paragraph("Opponents are tiered by CFP ranking (or BBMIF rank as fallback):")
    add_styled_table(doc,
        ["Opponent Tier", "Win", "Close Loss (<=3 pts)", "Loss"],
        [
            ["Top 10", "+6.0", "+1.5", "-1.5"],
            ["11-20", "+3.0", "+0.5", "-3.0"],
            ["21-25", "+3.0", "--", "-4.5"],
            ["26-40", "+1.0", "--", "--"],
            ["41-75", "+0.4", "--", "--"],
        ]
    )

    # Rank Caps
    doc.add_heading("Rank Caps", level=2)
    doc.add_paragraph("To prevent runaway rankings from limited data:")
    doc.add_paragraph("A team cannot rank more than 15 positions above their best-beaten opponent's rank", style='List Bullet')
    doc.add_paragraph("Day-over-day rank movement capped at 12 positions (or 6 for top-50 teams)", style='List Bullet')

    # Spread
    doc.add_heading("Spread Prediction", level=1)
    doc.add_heading("Formula", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "raw_line = (away_bbmif - home_bbmif + home_field_advantage) * 1.68\n"
        "final_line = snap_to_half_point(raw_line + 1.0)"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)

    doc.add_paragraph(
        "BBMIF_LINE_MULTIPLIER = 1.68 -- calibrated via OLS regression on 796 completed games (2025-26 season). "
        "Converts the BBMIF score differential into a point spread."
    )
    doc.add_paragraph(
        "HOME_BIAS_CORRECTION = +1.0 pts -- corrects observed systematic over-weighting of home teams. "
        "Our model picked home teams 70.5% of the time vs a 59.6% actual home win rate before this correction."
    )

    # HFA
    doc.add_heading("Home Field Advantage (Residual Method)", level=2)
    doc.add_paragraph(
        "Calculated per team using the residual method: for each completed game, compute "
        "residual = actual margin - projected margin. Average home residuals minus average away residuals, "
        "divided by 2. Scaled by 0.5 and capped at +/-8.0 points. This produces team-specific HFA values."
    )

    # Altitude
    doc.add_heading("Altitude Adjustments", level=2)
    add_styled_table(doc,
        ["Team", "Bonus (pts)"],
        [
            ["Air Force", "4.0"], ["Wyoming", "3.0"], ["Colorado / Colorado State", "2.5"],
            ["BYU / Utah", "2.0"], ["New Mexico", "1.5"], ["UNLV", "2.0"],
        ]
    )

    doc.add_paragraph("Bye week bonus: +2.5 points for teams coming off a bye vs non-bye opponents.")

    # O/U
    doc.add_heading("Over/Under Totals", level=1)
    doc.add_heading("Method", level=2)
    doc.add_paragraph("Point projections derived from SP+ ratings with 20% regression to mean:")
    p = doc.add_paragraph()
    run = p.add_run(
        "regressed_offense = raw_sp_offense * 0.80 + 29.0 * 0.20\n"
        "team_points = regressed_offense * (29.0 / opponent_regressed_defense)"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)

    doc.add_heading("Conference Tier Offsets", level=2)
    add_styled_table(doc,
        ["Matchup Type", "Total Adjustment"],
        [
            ["Power 4 vs Power 4", "-23.6 pts"],
            ["Power 4 vs Group of 5", "-12.6 pts"],
            ["Group of 5 vs Group of 5", "+4.7 pts"],
        ]
    )

    doc.add_heading("Weather Adjustments (outdoor games)", level=2)
    add_styled_table(doc,
        ["Condition", "Impact"],
        [
            ["Wind > 20 mph", "-4.0 pts"],
            ["Wind 15-20 mph", "-2.0 pts"],
            ["Temperature < 20F", "-3.0 pts"],
            ["Temperature 20-32F", "-1.5 pts"],
            ["Precipitation > 2mm", "-2.0 pts"],
        ]
    )

    doc.add_paragraph(
        "Pace adjustment: expected plays per game adjusts total by +/-20% relative to the 70-play FBS average."
    )

    # Win Prob
    doc.add_heading("Win Probability", level=1)
    doc.add_paragraph(
        "Gaussian CDF with the predicted spread as the mean and STD_DEV = 14.0 (calibrated via Brier score "
        "minimization). The actual observed standard deviation of game margins is 19.8, but 14.0 produces "
        "better-calibrated probabilities."
    )

    # Recency
    doc.add_heading("Recency Weighting", level=1)
    add_styled_table(doc,
        ["Parameter", "Value", "Purpose"],
        [
            ["Recency Decay", "0.85/week", "Exponential decay weighting recent box scores higher"],
            ["Early Season Penalty", "0.75x", "Reduce box score influence weeks 1-3"],
            ["Min Games Threshold", "7", "Fall back to season stats if fewer box scores available"],
        ]
    )

    # Data Pipeline
    doc.add_heading("Data Pipeline", level=1)
    add_styled_table(doc,
        ["Source", "Data", "Frequency"],
        [
            ["ESPN API", "Schedule, scores, detailed box scores", "Daily"],
            ["CFBD API", "SP+ ratings, team season stats", "Weekly (6-day cache)"],
            ["The Odds API", "Vegas spreads, totals, moneylines", "Each run"],
            ["Open-Meteo", "Temperature, wind speed, precipitation", "Cached per venue/date"],
        ]
    )

    # Calibration
    doc.add_heading("Calibration Infrastructure", level=1)
    doc.add_paragraph("The pipeline includes automated calibration modes:", style='List Bullet')
    calibrations = [
        "Weight Grid Search -- coordinate descent optimization over stat weights (~30-60 min)",
        "Line Multiplier Search -- OLS regression for BBMIF-to-spread conversion (<1 min)",
        "STD_DEV Search -- Brier score minimization for win probability calibration (<1 min)",
        "Total Offset Search -- conference-tier bias correction (<1 min)",
        "Rolling Backtest -- week-by-week point-in-time simulation for true predictive accuracy",
    ]
    for c in calibrations:
        doc.add_paragraph(c, style='List Bullet 2')

    # Known Issues
    doc.add_heading("Known Issues & Limitations", level=1)
    doc.add_paragraph("We are transparent about the model's current weaknesses:")

    issues = [
        ("SP+ Dependency (76% weight)", "The model is heavily reliant on a single external power rating. If SP+ is wrong about a team, BBMIF will be wrong."),
        ("Home Team Over-Selection", "Even after the 1.0 pt bias correction, the model still leans toward home teams more than Vegas does."),
        ("Mid-Season Accuracy Drop", "Weeks 1-3 ATS performance is strong (83% vs Vegas 73%), but weeks 4-9 drops to 52% vs Vegas 74%. Box score noise overwhelms the SP+ signal during the transition period."),
        ("No Injury Modeling", "Key quarterback absences can swing games 7-14 points. Not accounted for."),
        ("No Coaching Adjustments", "First-year coaches, coordinator changes, and scheme shifts are not modeled."),
        ("Flat Conference Tier Offsets", "Total offsets are static for the entire season."),
        ("No Motivation/Rivalry Factor", "Rivalry games, bowl incentives, and trap game dynamics not modeled."),
        ("Limited Pace Data", "Plays-per-game data not always available, so pace adjustments sometimes skipped."),
    ]
    for i, (title, desc) in enumerate(issues, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}: ")
        run.bold = True
        p.add_run(desc)

    # Questions
    doc.add_heading("What We're Looking For", level=1)
    doc.add_paragraph("We would appreciate your assessment of:")

    questions = [
        "Fundamental methodology -- Is the SP+-dominant weighted composite approach sound for college football, or should we use a different framework entirely?",
        "Spread calibration -- Is OLS regression the right approach for the line multiplier? Should we consider a different conversion method?",
        "Win probability -- Is Gaussian CDF with a fixed STD_DEV appropriate, or should we use a different distribution or variable STD_DEV?",
        "Over/under model -- The SP+ ratio method with static conference offsets feels like the weakest part. How do professional shops approach college football totals?",
        "Home field advantage -- Is the residual method reasonable, or is there a better approach? How should we handle neutral-site games?",
        "Injury impact -- What's the standard approach for quantifying point-spread impact of key player absences in college football?",
        "Weight calibration -- With SP+ at 76%, are we over-indexed on one input? What complementary signals would add the most predictive value?",
        "Seasonal dynamics -- How should the model handle the transition from preseason projections to in-season data? Our mid-season accuracy drop suggests the blending isn't right.",
        "Edge thresholds -- We currently recommend bets at 6.0+ point edge with spread < 14. Are these thresholds appropriate?",
        "Anything we're missing -- What factors do professional college football pricing models include that we've omitted entirely?",
    ]
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. {q}", style='List Number')

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("This document is confidential and intended solely for the purpose of the requested model review.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    return doc


if __name__ == "__main__":
    out_dir = os.path.dirname(os.path.abspath(__file__))

    doc1 = build_internal_doc()
    path1 = os.path.join(out_dir, "BBMI_Football_Model_Internal_Summary.docx")
    doc1.save(path1)
    print(f"Internal summary: {path1}")

    doc2 = build_review_doc()
    path2 = os.path.join(out_dir, "BBMI_Football_Model_Review_Brief.docx")
    doc2.save(path2)
    print(f"Review brief: {path2}")
